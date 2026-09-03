import streamlit as st
import time
from pathlib import Path
from typing import List, Dict
import pdfplumber
import pandas as pd
from openai import OpenAI
from docx import Document
import json
import io
import os


# ============================================================
# OpenRouter / OpenAI Client
# ============================================================

client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=os.getenv("OPENAI_API_KEY")
)


# ============================================================
# Read File
# ============================================================

def read_file(
    file_path: Path | io.BytesIO,
    file_extension: str
) -> str:

    try:

        if file_extension.lower() == ".txt":

            if isinstance(file_path, io.BytesIO):
                return file_path.read().decode("utf-8")

            return file_path.read_text(encoding="utf-8")

        elif file_extension.lower() == ".pdf":

            full_text = ""

            with pdfplumber.open(file_path) as pdf:

                for page_num, page in enumerate(pdf.pages):

                    page_text = page.extract_text()

                    if page_text:
                        full_text += page_text + "\n"

                    tables = page.extract_tables()

                    if tables:

                        for table in tables:

                            for row in table or []:

                                if not row or not any(
                                    cell and cell.strip()
                                    for cell in row
                                ):
                                    continue

                                row_text = " ".join(
                                    cell.strip().replace("\n", " ")
                                    if cell
                                    else ""
                                    for cell in row
                                )

                                if row_text.startswith(
                                    tuple(f"{i}." for i in range(1, 10))
                                ):
                                    full_text += f"\n{row_text}\n"

                                else:
                                    full_text += f" {row_text}\n"

            combined_text = full_text.strip().replace("\r", "")

            # Remove non-ASCII characters
            combined_text = (
                combined_text
                .encode("ascii", "ignore")
                .decode("ascii")
            )

            return combined_text

        elif file_extension.lower() == ".docx":

            doc = Document(file_path)

            return "\n".join(
                [
                    p.text
                    for p in doc.paragraphs
                    if p.text.strip()
                ]
            )

        else:

            raise ValueError(
                f"Unsupported file type: {file_extension}"
            )

    except Exception as e:

        st.error(f"Failed to read file: {e}")

        return ""


# ============================================================
# Clean Text
# ============================================================

def clean_text(text: str) -> str:

    if not text:
        return ""

    text = text.replace("–", "-")
    text = text.replace("Page ", "").replace(" of ", "")

    lines = text.split("\n")

    cleaned_lines = [
        line.strip()
        for line in lines
        if line.strip()
    ]

    return "\n".join(cleaned_lines)


# ============================================================
# Chunk PDF
# ============================================================

def chunk_pdf(
    file_path: Path | io.BytesIO,
    max_pages=20
) -> List[str]:

    chunks = []

    try:

        with pdfplumber.open(file_path) as pdf:

            total_pages = len(pdf.pages)

            for i in range(0, total_pages, max_pages):

                chunk_text = ""

                for page in pdf.pages[i:i + max_pages]:

                    page_text = page.extract_text() or ""

                    chunk_text += page_text + "\n"

                cleaned_chunk = clean_text(chunk_text)

                if cleaned_chunk:
                    chunks.append(cleaned_chunk)

    except Exception as e:

        st.error(f"Error chunking PDF: {e}")

    return chunks


# ============================================================
# Chunk DOCX
# ============================================================

def chunk_docx(
    file_path: Path | io.BytesIO,
    max_words=5000
) -> List[str]:

    chunks = []

    current = []

    word_count = 0

    try:

        doc = Document(file_path)

        for para in doc.paragraphs:

            words = para.text.split()

            if word_count + len(words) > max_words:

                if current:

                    chunks.append(
                        clean_text(" ".join(current))
                    )

                current = []

                word_count = 0

            current.extend(words)

            word_count += len(words)

        if current:

            chunks.append(
                clean_text(" ".join(current))
            )

    except Exception as e:

        st.error(f"Error chunking DOCX: {e}")

    return chunks


# ============================================================
# Extract JSON From Model Response
# ============================================================

def extract_json(content: str) -> str:

    if not content:
        return ""

    brace_count = 0

    start_idx = None

    json_content = ""

    for i, char in enumerate(content):

        if char == "{":

            if brace_count == 0:
                start_idx = i

            brace_count += 1

        elif char == "}":

            if brace_count > 0:
                brace_count -= 1

            if (
                brace_count == 0
                and start_idx is not None
            ):

                json_content = content[
                    start_idx:i + 1
                ]

                break

    return json_content or content


# ============================================================
# Extract Tasks
# ============================================================

def extract_task(
    text: str
) -> Dict[str, List[Dict[str, str]]]:

    prompt = f"""
You are an expert federal proposal analyst and
technical project planner.

Carefully review the solicitation document and produce
structured JSON output.

First, identify the major task headings exactly as they
appear in the document, including section headings,
task titles, numbered tasks, or major work areas.

Use these exact task names as "Parent Task" values.

Do NOT invent or infer major task names that are not
present in the source document.

For each subtask under these major tasks, provide:

"Task": Subtask name or description, focusing on
technical tasks and compliance requirements.

"Parent Task": One of the major tasks identified,
using the exact wording from the document.

"Methodology": Methodology for accomplishing the task.
If the solicitation does not specify a methodology,
use "Agile (ADLC) / Secure SDLC".

"Tools & Technologies": Tools, platforms, standards,
technologies, frameworks, and compliance standards
explicitly identified or reasonably required by the
solicitation.

"Task Summary": 2-3 sentences explaining the scope,
objective, and expected outcome of the task.

For project management deliverables such as PMP,
Integrated Master Schedule, Risk Management Plan,
Configuration Management Plan, Quality Assurance Plan,
etc., create a separate "Deliverables" list.

Each deliverable must contain:

"Deliverable": Name or description.

"Parent Task": Associated major task, using the exact
wording from the document.

"Description": Brief explanation of the deliverable.

IMPORTANT RULES:

Do not invent requirements.

Do not create technologies that are not supported by
the solicitation.

Preserve exact solicitation terminology for task and
parent-task names.

Separate actual solicitation requirements from
methodology recommendations.

Capture compliance requirements when explicitly stated.

Capture technical standards, tools, platforms, and
technologies when stated.

If a requirement is unclear, preserve the source
language instead of guessing.

Return ONLY valid JSON.

Do not include Markdown code fences.

Do not include explanations outside the JSON object.

Return exactly this structure:

{{
    "Tasks": [
        {{
            "Task": "...",
            "Parent Task": "...",
            "Methodology": "...",
            "Tools & Technologies": "...",
            "Task Summary": "..."
        }}
    ],
    "Deliverables": [
        {{
            "Deliverable": "...",
            "Parent Task": "...",
            "Description": "..."
        }}
    ]
}}

Document:

{text}
"""

    try:

        response = client.chat.completions.create(

            model="liquid/lfm-2.5-2.6b:free",

            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],

            temperature=0.2
        )

        # Safely retrieve model content
        message = response.choices[0].message

        content = message.content if message else ""

        if not content:

            st.warning(
                "The AI model returned an empty response."
            )

            return {
                "Tasks": [],
                "Deliverables": []
            }

        content = content.strip()

        json_content = extract_json(content)

        try:

            result = json.loads(json_content)

            if not isinstance(result, dict):

                raise ValueError(
                    "AI response was not a JSON object."
                )

            # Ensure expected keys exist
            if "Tasks" not in result:
                result["Tasks"] = []

            if "Deliverables" not in result:
                result["Deliverables"] = []

            # Remove unspecified parent tasks
            result["Tasks"] = [
                task
                for task in result.get("Tasks", [])
                if isinstance(task, dict)
                and task.get("Parent Task")
                != "Unspecified Task"
            ]

            result["Deliverables"] = [
                deliverable
                for deliverable in result.get(
                    "Deliverables",
                    []
                )
                if isinstance(deliverable, dict)
                and deliverable.get("Parent Task")
                != "Unspecified Task"
            ]

            return result

        except json.JSONDecodeError as json_error:

            st.error(
                "Could not parse JSON response from AI model: "
                f"{json_error}"
            )

            st.write(
                "Failed JSON Content:"
            )

            st.code(
                json_content,
                language="json"
            )

            return {
                "Tasks": [],
                "Deliverables": []
            }

    except Exception as e:

        st.error(
            f"Error during AI extraction: {e}"
        )

        return {
            "Tasks": [],
            "Deliverables": []
        }


# ============================================================
# Process File
# ============================================================

def process_file(
    file_path: Path | io.BytesIO,
    file_extension: str
) -> Dict[str, List[Dict[str, str]]]:

    extracted_results = {
        "Tasks": [],
        "Deliverables": []
    }

    # --------------------------------------------------------
    # Read and clean complete document
    # --------------------------------------------------------

    full_text = clean_text(
        read_file(
            file_path,
            file_extension
        )
    )

    if not full_text:

        st.warning(
            "No text could be extracted from this file."
        )

        return extracted_results

    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    if file_extension.lower() == ".pdf":

        try:

            with pdfplumber.open(file_path) as pdf:

                total_pages = len(pdf.pages)

        except Exception as e:

            st.error(
                f"Could not determine PDF page count: {e}"
            )

            return extracted_results

        if total_pages > 40:

            st.info(
                f"Splitting PDF into chunks "
                f"({total_pages} pages)..."
            )

            chunks = chunk_pdf(file_path)

        else:

            chunks = [full_text]

    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    elif file_extension.lower() == ".docx":

        try:

            doc = Document(file_path)

            total_words = sum(
                len(p.text.split())
                for p in doc.paragraphs
            )

        except Exception as e:

            st.error(
                f"Could not read DOCX: {e}"
            )

            return extracted_results

        if total_words > 20000:

            st.info(
                f"Splitting DOCX into chunks "
                f"(~{total_words} words)..."
            )

            chunks = chunk_docx(file_path)

        else:

            chunks = [full_text]

    # --------------------------------------------------------
    # TXT
    # --------------------------------------------------------

    else:

        chunks = [full_text]

    # Remove empty chunks
    chunks = [
        chunk
        for chunk in chunks
        if chunk and chunk.strip()
    ]

    if not chunks:

        st.warning(
            "No usable text chunks were created."
        )

        return extracted_results

    # --------------------------------------------------------
    # Process Chunks
    # --------------------------------------------------------

    for i, chunk in enumerate(chunks, 1):

        st.info(
            f"Extracting chunk "
            f"{i}/{len(chunks)}..."
        )

        chunk_result = extract_task(chunk)

        if not isinstance(chunk_result, dict):

            continue

        extracted_results["Tasks"].extend(
            chunk_result.get("Tasks", [])
        )

        extracted_results["Deliverables"].extend(
            chunk_result.get("Deliverables", [])
        )

    # --------------------------------------------------------
    # Consolidate Tasks
    # --------------------------------------------------------

    consolidated_tasks = {}

    for task in extracted_results["Tasks"]:

        if not isinstance(task, dict):
            continue

        task_key = (
            task.get("Task", "").strip(),
            task.get("Parent Task", "").strip()
        )

        if task_key not in consolidated_tasks:

            consolidated_tasks[task_key] = task

        else:

            existing_summary = (
                consolidated_tasks[
                    task_key
                ].get(
                    "Task Summary",
                    ""
                )
            )

            new_summary = task.get(
                "Task Summary",
                ""
            )

            if new_summary:

                consolidated_tasks[
                    task_key
                ]["Task Summary"] = (
                    existing_summary
                    + " "
                    + new_summary
                ).strip()

    # --------------------------------------------------------
    # Consolidate Deliverables
    # --------------------------------------------------------

    consolidated_deliverables = {}

    for deliverable in extracted_results[
        "Deliverables"
    ]:

        if not isinstance(deliverable, dict):
            continue

        deliv_key = (
            deliverable.get(
                "Deliverable",
                ""
            ).strip(),

            deliverable.get(
                "Parent Task",
                ""
            ).strip()
        )

        if deliv_key not in consolidated_deliverables:

            consolidated_deliverables[
                deliv_key
            ] = deliverable

        else:

            existing_description = (
                consolidated_deliverables[
                    deliv_key
                ].get(
                    "Description",
                    ""
                )
            )

            new_description = deliverable.get(
                "Description",
                ""
            )

            if new_description:

                consolidated_deliverables[
                    deliv_key
                ]["Description"] = (
                    existing_description
                    + " "
                    + new_description
                ).strip()

    # --------------------------------------------------------
    # Return Consolidated Results
    # --------------------------------------------------------

    extracted_results["Tasks"] = list(
        consolidated_tasks.values()
    )

    extracted_results["Deliverables"] = list(
        consolidated_deliverables.values()
    )

    return extracted_results


# ============================================================
# Streamlit UI
# ============================================================

st.title(
    "Sara: Software Automation for Requirement Analysis"
)

st.write(
    "Upload one or more solicitation documents "
    "(TXT, PDF, or DOCX) to extract tasks and deliverables."
)


# ============================================================
# File Uploader
# ============================================================

uploaded_files = st.file_uploader(
    "Choose files",
    type=["txt", "pdf", "docx"],
    accept_multiple_files=True
)


# ============================================================
# Process Uploaded Files
# ============================================================

if uploaded_files:

    temp_dir = Path("temp")

    temp_dir.mkdir(exist_ok=True)

    all_tasks = []

    all_deliverables = []

    start_time = time.time()

    # --------------------------------------------------------
    # Process Each File
    # --------------------------------------------------------

    for uploaded_file in uploaded_files:

        st.write(
            f"Processing {uploaded_file.name}..."
        )

        file_extension = (
            f".{uploaded_file.name.split('.')[-1].lower()}"
        )

        temp_file_path = (
            temp_dir / uploaded_file.name
        )

        # ----------------------------------------------------
        # Save Uploaded File Temporarily
        # ----------------------------------------------------

        try:

            with open(
                temp_file_path,
                "wb"
            ) as f:

                f.write(
                    uploaded_file.getbuffer()
                )

        except Exception as e:

            st.error(
                f"Could not save "
                f"{uploaded_file.name}: {e}"
            )

            continue

        # ----------------------------------------------------
        # Process File
        # ----------------------------------------------------

        try:

            extracted = process_file(
                temp_file_path,
                file_extension
            )

        except Exception as e:

            st.error(
                f"Error processing "
                f"{uploaded_file.name}: {e}"
            )

            extracted = {
                "Tasks": [],
                "Deliverables": []
            }

        # ----------------------------------------------------
        # Add Source File
        # ----------------------------------------------------

        for task in extracted["Tasks"]:

            task["Source File"] = (
                uploaded_file.name
            )

        for deliverable in extracted[
            "Deliverables"
        ]:

            deliverable["Source File"] = (
                uploaded_file.name
            )

        # ----------------------------------------------------
        # Append to Aggregated Results
        # ----------------------------------------------------

        all_tasks.extend(
            extracted["Tasks"]
        )

        all_deliverables.extend(
            extracted["Deliverables"]
        )

        # ====================================================
        # Display Results for Current File
        # ====================================================

        st.subheader(
            f"Results for {uploaded_file.name}"
        )

        # ----------------------------------------------------
        # Display Tasks
        # ----------------------------------------------------

        if extracted["Tasks"]:

            st.write(
                "**Extracted Tasks**"
            )

            tasks_df = pd.DataFrame(
                extracted["Tasks"]
            )

            st.dataframe(
                tasks_df,
                use_container_width=True
            )

            # Excel download
            tasks_excel = io.BytesIO()

            tasks_df.to_excel(
                tasks_excel,
                index=False
            )

            tasks_excel.seek(0)

            st.download_button(
                label=(
                    f"Download Tasks for "
                    f"{uploaded_file.name} "
                    f"as Excel"
                ),

                data=tasks_excel,

                file_name=(
                    f"{Path(uploaded_file.name).stem}"
                    "_tasks.xlsx"
                ),

                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.spreadsheetml.sheet"
                )
            )

        else:

            st.warning(
                f"No tasks extracted from "
                f"{uploaded_file.name}."
            )

        # ----------------------------------------------------
        # Display Deliverables
        # ----------------------------------------------------

        if extracted["Deliverables"]:

            st.write(
                "**Extracted Deliverables**"
            )

            deliverables_df = pd.DataFrame(
                extracted["Deliverables"]
            )

            st.dataframe(
                deliverables_df,
                use_container_width=True
            )

            # Excel download
            deliverables_excel = io.BytesIO()

            deliverables_df.to_excel(
                deliverables_excel,
                index=False
            )

            deliverables_excel.seek(0)

            st.download_button(
                label=(
                    f"Download Deliverables for "
                    f"{uploaded_file.name} "
                    f"as Excel"
                ),

                data=deliverables_excel,

                file_name=(
                    f"{Path(uploaded_file.name).stem}"
                    "_deliverables.xlsx"
                ),

                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.spreadsheetml.sheet"
                )
            )

        else:

            st.warning(
                f"No deliverables extracted from "
                f"{uploaded_file.name}."
            )

        # ----------------------------------------------------
        # Clean Up Temporary File
        # ----------------------------------------------------

        try:

            os.remove(temp_file_path)

        except OSError:

            pass

    # ========================================================
    # Aggregated Results
    # ========================================================

    if all_tasks or all_deliverables:

        st.subheader(
            "Aggregated Results Across All Files"
        )

        # ----------------------------------------------------
        # All Tasks
        # ----------------------------------------------------

        if all_tasks:

            st.write(
                "**All Extracted Tasks**"
            )

            all_tasks_df = pd.DataFrame(
                all_tasks
            )

            st.dataframe(
                all_tasks_df,
                use_container_width=True
            )

            tasks_excel = io.BytesIO()

            all_tasks_df.to_excel(
                tasks_excel,
                index=False
            )

            tasks_excel.seek(0)

            st.download_button(
                label="Download All Tasks as Excel",

                data=tasks_excel,

                file_name="all_tasks.xlsx",

                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.spreadsheetml.sheet"
                )
            )

        # ----------------------------------------------------
        # All Deliverables
        # ----------------------------------------------------

        if all_deliverables:

            st.write(
                "**All Extracted Deliverables**"
            )

            all_deliverables_df = pd.DataFrame(
                all_deliverables
            )

            st.dataframe(
                all_deliverables_df,
                use_container_width=True
            )

            deliverables_excel = io.BytesIO()

            all_deliverables_df.to_excel(
                deliverables_excel,
                index=False
            )

            deliverables_excel.seek(0)

            st.download_button(
                label=(
                    "Download All Deliverables "
                    "as Excel"
                ),

                data=deliverables_excel,

                file_name="all_deliverables.xlsx",

                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.spreadsheetml.sheet"
                )
            )

    # ========================================================
    # Processing Time
    # ========================================================

    end_time = time.time()

    elapsed = round(
        end_time - start_time,
        2
    )

    st.success(
        f"Finished processing "
        f"{len(uploaded_files)} file(s) "
        f"in {elapsed} seconds."
    )

else:

    st.info(
        "Please upload one or more files "
        "to begin processing."
    )

